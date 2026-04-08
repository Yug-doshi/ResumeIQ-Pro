"""
File Handler — Extract text from uploaded resume files (PDF and DOCX)

Supports:
  - PDF extraction via PyPDF2
  - DOCX extraction via python-docx
  - Fallback text cleaning
"""

import os


def extract_text_from_resume(file_path: str) -> str:
    """
    Extract text content from a resume file.

    Args:
        file_path: Path to the uploaded resume (PDF or DOCX)

    Returns:
        Extracted plain text string
    """
    file_ext = os.path.splitext(file_path)[1].lower()

    if file_ext == ".pdf":
        return _extract_from_pdf(file_path)
    elif file_ext == ".docx":
        return _extract_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_ext}")


def _extract_from_pdf(file_path: str) -> str:
    """Extract text from a PDF file using PyPDF2."""
    try:
        from PyPDF2 import PdfReader

        reader = PdfReader(file_path)
        text_parts = []

        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

        full_text = "\n".join(text_parts)

        if not full_text.strip():
            return "Unable to extract text from this PDF. It may be image-based. Please try a text-based PDF."

        return _clean_extracted_text(full_text)

    except Exception as e:
        raise ValueError(f"Failed to read PDF: {str(e)}")


def _extract_from_docx(file_path: str) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        from docx import Document

        doc = Document(file_path)
        text_parts = []

        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Extract from tables too
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)

        full_text = "\n".join(text_parts)

        if not full_text.strip():
            return "Unable to extract text from this DOCX file."

        return _clean_extracted_text(full_text)

    except Exception as e:
        raise ValueError(f"Failed to read DOCX: {str(e)}")


def _clean_extracted_text(text: str) -> str:
    """Clean up extracted text: remove excess whitespace, normalize line breaks."""
    # Replace multiple newlines with double newline
    import re
    text = re.sub(r'\n{3,}', '\n\n', text)
    # Replace multiple spaces with single space
    text = re.sub(r'[ \t]+', ' ', text)
    # Strip each line
    lines = [line.strip() for line in text.split('\n')]
    text = '\n'.join(lines)
    return text.strip()
